"""
NHS-Specific PII Redactor with GUI
Combines Presidio's ML-powered detection with NHS-specific patterns
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import threading
import sys
import re
from pathlib import Path
from typing import List, Dict, Optional
from datetime import datetime

try:
    from presidio_analyzer import AnalyzerEngine, Pattern, PatternRecognizer
    from presidio_analyzer.nlp_engine import NlpEngineProvider
    from presidio_anonymizer import AnonymizerEngine
    from presidio_anonymizer.entities import OperatorConfig
except ImportError:
    print("Please install Presidio:")
    print("pip install presidio-analyzer presidio-anonymizer")
    print("pip install spacy")
    print("python -m spacy download en_core_web_lg")
    sys.exit(1)

try:
    from docx import Document
except ImportError:
    print("Please install python-docx: pip install python-docx")
    sys.exit(1)

try:
    import fitz
except ImportError:
    print("Please install PyMuPDF: pip install PyMuPDF")
    sys.exit(1)


class NHSPatternRecognizer(PatternRecognizer):
    """Custom recognizer for NHS-specific PII patterns"""
    
    def __init__(self):
        patterns = [
            Pattern(
                name="nhs_number_pattern",
                regex=r"\b\d{3}[\s-]?\d{3}[\s-]?\d{4}\b",
                score=0.85,
            ),
            Pattern(
                name="hospital_number_pattern",
                regex=r"\b(?:Hospital|Patient|Trust|MRN|Hosp)[\s#:]*([A-Z0-9]{5,12})\b",
                score=0.8,
            ),
            Pattern(
                name="chi_number_pattern",
                regex=r"\bCHI[\s:]*\d{10}\b",
                score=0.9,
            ),
            Pattern(
                name="gmc_number_pattern",
                regex=r"\bGMC[\s:#]*\d{7}\b",
                score=0.9,
            ),
            Pattern(
                name="gp_practice_code_pattern",
                regex=r"\b(?:GP|Practice|Surgery)[\s:]+([A-Z]\d{5})\b",
                score=0.75,
            ),
            Pattern(
                name="uk_driving_license_pattern",
                regex=r"\b[A-Z]{5}\d{6}[A-Z]{2}\d[A-Z]{2}\b",
                score=0.85,
            ),
        ]
        
        super().__init__(
            supported_entity="NHS_IDENTIFIER",
            patterns=patterns,
            context=["NHS", "patient", "hospital", "medical", "clinic", "GP"],
        )


class UKPhoneRecognizer(PatternRecognizer):
    """Custom recognizer for UK phone numbers"""
    
    def __init__(self):
        patterns = [
            Pattern(
                name="uk_mobile_pattern",
                regex=r"\b(?:07\d{3}|447\d{3})[\s-]?\d{6}\b",
                score=0.85,
            ),
            Pattern(
                name="uk_landline_pattern",
                regex=r"\b(?:0|\+44)[\s-]?\d{3,4}[\s-]?\d{3,4}[\s-]?\d{3,4}\b",
                score=0.7,
            ),
        ]
        
        super().__init__(
            supported_entity="UK_PHONE_NUMBER",
            patterns=patterns,
            context=["phone", "mobile", "tel", "telephone", "contact"],
        )


class UKPostcodeRecognizer(PatternRecognizer):
    """Custom recognizer for UK postcodes"""
    
    def __init__(self):
        patterns = [
            Pattern(
                name="uk_postcode_pattern",
                regex=r"\b[A-Z]{1,2}\d{1,2}[A-Z]?\s?\d[A-Z]{2}\b",
                score=0.85,
            ),
        ]
        
        super().__init__(
            supported_entity="UK_POSTCODE",
            patterns=patterns,
            context=["address", "postcode", "postal", "location"],
        )


class UKNINumberRecognizer(PatternRecognizer):
    """Custom recognizer for UK National Insurance Numbers"""
    
    def __init__(self):
        patterns = [
            Pattern(
                name="ni_number_pattern",
                regex=r"\b[A-CEGHJ-PR-TW-Z]{2}[\s-]?\d{2}[\s-]?\d{2}[\s-]?\d{2}[\s-]?[A-D]\b",
                score=0.9,
            ),
        ]
        
        super().__init__(
            supported_entity="UK_NI_NUMBER",
            patterns=patterns,
            context=["NI", "National Insurance", "insurance number"],
        )


class NHSPresidioRedactor:
    """NHS-specific PII redactor using Presidio with custom patterns"""
    
    def __init__(self, log_callback=None):
        """
        Initialize the NHS Presidio redactor.
        
        Args:
            log_callback: Function to call for logging messages
        """
        self.log_callback = log_callback
        self.analyzer = None
        self.anonymizer = None
        
        self.entities = [
            "PERSON", "EMAIL_ADDRESS", "PHONE_NUMBER", "DATE_TIME",
            "LOCATION", "NHS_IDENTIFIER", "UK_PHONE_NUMBER",
            "UK_POSTCODE", "UK_NI_NUMBER", "MEDICAL_LICENSE", "UK_NHS",
        ]
    
    def log(self, message):
        """Log a message"""
        if self.log_callback:
            self.log_callback(message)
        else:
            print(message)
    
    def initialize(self):
        """Initialize Presidio (can be slow, so called separately)"""
        self.log("🔧 Initializing Presidio analyzer...")
        self.log("This may take 30-60 seconds on first run...")
        
        try:
            configuration = {
                "nlp_engine_name": "spacy",
                "models": [{"lang_code": "en", "model_name": "en_core_web_lg"}],
            }
            
            provider = NlpEngineProvider(nlp_configuration=configuration)
            nlp_engine = provider.create_engine()
            
            self.analyzer = AnalyzerEngine(nlp_engine=nlp_engine)
            self.analyzer.registry.add_recognizer(NHSPatternRecognizer())
            self.analyzer.registry.add_recognizer(UKPhoneRecognizer())
            self.analyzer.registry.add_recognizer(UKPostcodeRecognizer())
            self.analyzer.registry.add_recognizer(UKNINumberRecognizer())
            
            self.anonymizer = AnonymizerEngine()
            
            self.log("✅ Presidio analyzer ready!")
            return True
        except Exception as e:
            self.log(f"❌ Error initializing Presidio: {e}")
            return False
    
    def analyze_text(self, text: str, language: str = "en") -> List:
        """Analyze text and return detected PII entities"""
        if not self.analyzer:
            raise RuntimeError("Analyzer not initialized. Call initialize() first.")
        
        results = self.analyzer.analyze(
            text=text,
            language=language,
            entities=self.entities,
            allow_list=None,
            score_threshold=0.3,
        )
        return results
    
    def anonymize_text(self, text: str, language: str = "en") -> tuple:
        """Anonymize text by replacing PII with labels"""
        results = self.analyze_text(text, language)
        
        findings = []
        for result in results:
            findings.append({
                'type': result.entity_type,
                'start': result.start,
                'end': result.end,
                'score': result.score,
                'text': text[result.start:result.end]
            })
        
        anonymized = self.anonymizer.anonymize(
            text=text,
            analyzer_results=results,
            operators={
                "DEFAULT": OperatorConfig("replace", {"new_value": "[REDACTED]"}),
                "PERSON": OperatorConfig("replace", {"new_value": "[REDACTED-NAME]"}),
                "EMAIL_ADDRESS": OperatorConfig("replace", {"new_value": "[REDACTED-EMAIL]"}),
                "PHONE_NUMBER": OperatorConfig("replace", {"new_value": "[REDACTED-PHONE]"}),
                "UK_PHONE_NUMBER": OperatorConfig("replace", {"new_value": "[REDACTED-UK-PHONE]"}),
                "UK_POSTCODE": OperatorConfig("replace", {"new_value": "[REDACTED-POSTCODE]"}),
                "UK_NI_NUMBER": OperatorConfig("replace", {"new_value": "[REDACTED-NI-NUMBER]"}),
                "NHS_IDENTIFIER": OperatorConfig("replace", {"new_value": "[REDACTED-NHS-ID]"}),
                "UK_NHS": OperatorConfig("replace", {"new_value": "[REDACTED-NHS-NUMBER]"}),
                "DATE_TIME": OperatorConfig("replace", {"new_value": "[REDACTED-DATE]"}),
                "LOCATION": OperatorConfig("replace", {"new_value": "[REDACTED-LOCATION]"}),
                "MEDICAL_LICENSE": OperatorConfig("replace", {"new_value": "[REDACTED-MEDICAL-LICENSE]"}),
            }
        )
        
        return anonymized.text, len(results), findings
    
    def redact_docx(self, input_path: str, output_path: str) -> Dict:
        """Redact PII from Word document"""
        self.log("📄 Loading Word document...")
        doc = Document(input_path)
        
        stats = {
            'paragraphs': 0,
            'tables': 0,
            'headers': 0,
            'footers': 0,
            'total_redactions': 0,
            'by_type': {}
        }
        
        self.log("🔍 Scanning and redacting paragraphs...")
        for para in doc.paragraphs:
            if para.text.strip():
                anonymized, count, findings = self.anonymize_text(para.text)
                if count > 0:
                    para.text = anonymized
                    stats['paragraphs'] += count
                    self._update_type_stats(stats, findings)
        
        self.log("🔍 Scanning and redacting tables...")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        anonymized, count, findings = self.anonymize_text(cell.text)
                        if count > 0:
                            cell.text = anonymized
                            stats['tables'] += count
                            self._update_type_stats(stats, findings)
        
        self.log("🔍 Scanning and redacting headers/footers...")
        for section in doc.sections:
            for header in section.header.paragraphs:
                if header.text.strip():
                    anonymized, count, findings = self.anonymize_text(header.text)
                    if count > 0:
                        header.text = anonymized
                        stats['headers'] += count
                        self._update_type_stats(stats, findings)
            
            for footer in section.footer.paragraphs:
                if footer.text.strip():
                    anonymized, count, findings = self.anonymize_text(footer.text)
                    if count > 0:
                        footer.text = anonymized
                        stats['footers'] += count
                        self._update_type_stats(stats, findings)
        
        stats['total_redactions'] = (stats['paragraphs'] + stats['tables'] + 
                                     stats['headers'] + stats['footers'])
        
        self.log("💾 Saving redacted document...")
        doc.save(output_path)
        
        return stats
    
    def redact_pdf(self, input_path: str, output_path: str) -> Dict:
        """Redact PII from PDF document"""
        self.log("📄 Loading PDF document...")
        doc = fitz.open(input_path)
        
        stats = {
            'pages_processed': 0,
            'total_redactions': 0,
            'by_type': {}
        }
        
        total_pages = len(doc)
        self.log(f"🔍 Scanning and redacting {total_pages} pages...")
        
        for page_num, page in enumerate(doc, 1):
            text = page.get_text()
            results = self.analyze_text(text)
            
            if results:
                for result in results:
                    pii_text = text[result.start:result.end]
                    areas = page.search_for(pii_text)
                    
                    if areas:
                        for area in areas:
                            page.add_redact_annot(area, fill=(0, 0, 0))
                            stats['total_redactions'] += 1
                            
                            entity_type = result.entity_type
                            stats['by_type'][entity_type] = stats['by_type'].get(entity_type, 0) + 1
                
                page.apply_redactions()
            
            stats['pages_processed'] += 1
            
            if page_num % 5 == 0:
                self.log(f"   Processed {page_num}/{total_pages} pages...")
        
        self.log("💾 Saving redacted PDF...")
        doc.save(output_path)
        doc.close()
        
        return stats
    
    def _update_type_stats(self, stats: Dict, findings: List[Dict]):
        """Update statistics by PII type"""
        for finding in findings:
            entity_type = finding['type']
            stats['by_type'][entity_type] = stats['by_type'].get(entity_type, 0) + 1
    
    def verify_redaction(self, file_path: str, file_type: str) -> Dict:
        """Verify that PII has been removed from the redacted file"""
        self.log("🔍 Verifying redaction...")
        verification = {
            'pii_found': 0,
            'types_found': {},
            'safe': True
        }
        
        try:
            if file_type == '.docx':
                doc = Document(file_path)
                full_text = '\n'.join([para.text for para in doc.paragraphs])
                
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            full_text += '\n' + cell.text
            
            elif file_type == '.pdf':
                doc = fitz.open(file_path)
                full_text = ''
                for page in doc:
                    full_text += page.get_text()
                doc.close()
            
            else:
                return verification
            
            results = self.analyze_text(full_text)
            
            if results:
                verification['safe'] = False
                verification['pii_found'] = len(results)
                
                for result in results:
                    entity_type = result.entity_type
                    verification['types_found'][entity_type] = \
                        verification['types_found'].get(entity_type, 0) + 1
        
        except Exception as e:
            verification['error'] = str(e)
        
        return verification


class NHSRedactorGUI:
    """GUI Application for NHS PII Redactor"""
    
    def __init__(self, root):
        self.root = root
        self.root.title("NHS PII Redactor - Presidio Edition")
        self.root.geometry("800x600")
        
        self.redactor = None
        self.input_file = None
        self.output_file = None
        
        self.setup_ui()
        
        # Initialize Presidio in background
        self.log_message("Welcome to NHS PII Redactor!")
        self.log_message("Initializing Presidio analyzer...")
        threading.Thread(target=self.initialize_presidio, daemon=True).start()
    
    def setup_ui(self):
        """Setup the user interface"""
        
        # Title
        title_frame = ttk.Frame(self.root, padding="10")
        title_frame.pack(fill=tk.X)
        
        title_label = ttk.Label(
            title_frame,
            text="NHS PII Redactor",
            font=("Arial", 16, "bold")
        )
        title_label.pack()
        
        subtitle_label = ttk.Label(
            title_frame,
            text="Powered by Microsoft Presidio + NHS-Specific Patterns",
            font=("Arial", 9)
        )
        subtitle_label.pack()
        
        # File selection frame
        file_frame = ttk.LabelFrame(self.root, text="File Selection", padding="10")
        file_frame.pack(fill=tk.X, padx=10, pady=5)
        
        # Input file
        input_frame = ttk.Frame(file_frame)
        input_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(input_frame, text="Input File:", width=12).pack(side=tk.LEFT)
        self.input_entry = ttk.Entry(input_frame, width=50)
        self.input_entry.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        ttk.Button(
            input_frame,
            text="Browse...",
            command=self.browse_input
        ).pack(side=tk.LEFT)
        
        # Output file
        output_frame = ttk.Frame(file_frame)
        output_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(output_frame, text="Output File:", width=12).pack(side=tk.LEFT)
        self.output_entry = ttk.Entry(output_frame, width=50)
        self.output_entry.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        ttk.Button(
            output_frame,
            text="Browse...",
            command=self.browse_output
        ).pack(side=tk.LEFT)
        
        # Action buttons
        button_frame = ttk.Frame(self.root, padding="10")
        button_frame.pack(fill=tk.X)
        
        self.redact_button = ttk.Button(
            button_frame,
            text="🔒 Redact Document",
            command=self.redact_document,
            state=tk.DISABLED
        )
        self.redact_button.pack(side=tk.LEFT, padx=5)
        
        ttk.Button(
            button_frame,
            text="Clear",
            command=self.clear_form
        ).pack(side=tk.LEFT, padx=5)
        
        # Progress frame
        progress_frame = ttk.Frame(self.root, padding="10")
        progress_frame.pack(fill=tk.X)
        
        self.progress = ttk.Progressbar(
            progress_frame,
            mode='indeterminate'
        )
        self.progress.pack(fill=tk.X)
        
        # Log output
        log_frame = ttk.LabelFrame(self.root, text="Processing Log", padding="10")
        log_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        self.log_text = scrolledtext.ScrolledText(
            log_frame,
            height=15,
            wrap=tk.WORD,
            font=("Consolas", 9)
        )
        self.log_text.pack(fill=tk.BOTH, expand=True)
        
        # Status bar
        self.status_var = tk.StringVar(value="Initializing...")
        status_bar = ttk.Label(
            self.root,
            textvariable=self.status_var,
            relief=tk.SUNKEN,
            anchor=tk.W
        )
        status_bar.pack(fill=tk.X, side=tk.BOTTOM)
    
    def log_message(self, message):
        """Add a message to the log"""
        self.log_text.insert(tk.END, message + "\n")
        self.log_text.see(tk.END)
        self.root.update_idletasks()
    
    def initialize_presidio(self):
        """Initialize Presidio in background"""
        try:
            self.redactor = NHSPresidioRedactor(log_callback=self.log_message)
            success = self.redactor.initialize()
            
            if success:
                self.root.after(0, self.enable_redaction)
                self.root.after(0, lambda: self.status_var.set("Ready"))
            else:
                self.root.after(0, lambda: self.status_var.set("Initialization failed"))
                self.root.after(0, lambda: messagebox.showerror(
                    "Error",
                    "Failed to initialize Presidio. Please check the log for details."
                ))
        except Exception as e:
            self.log_message(f"❌ Error: {e}")
            self.root.after(0, lambda: self.status_var.set("Error"))
    
    def enable_redaction(self):
        """Enable the redact button"""
        self.redact_button.config(state=tk.NORMAL)
    
    def browse_input(self):
        """Browse for input file"""
        filename = filedialog.askopenfilename(
            title="Select Document to Redact",
            filetypes=[
                ("Word Documents", "*.docx"),
                ("PDF Files", "*.pdf"),
                ("All Files", "*.*")
            ]
        )
        
        if filename:
            self.input_file = filename
            self.input_entry.delete(0, tk.END)
            self.input_entry.insert(0, filename)
            
            # Auto-suggest output filename
            input_path = Path(filename)
            output_path = input_path.parent / f"{input_path.stem}_REDACTED{input_path.suffix}"
            self.output_entry.delete(0, tk.END)
            self.output_entry.insert(0, str(output_path))
    
    def browse_output(self):
        """Browse for output file"""
        if self.input_file:
            input_path = Path(self.input_file)
            default_ext = input_path.suffix
            
            if default_ext == '.docx':
                filetypes = [("Word Documents", "*.docx")]
            elif default_ext == '.pdf':
                filetypes = [("PDF Files", "*.pdf")]
            else:
                filetypes = [("All Files", "*.*")]
        else:
            filetypes = [("All Files", "*.*")]
        
        filename = filedialog.asksaveasfilename(
            title="Save Redacted Document As",
            filetypes=filetypes,
            defaultextension=default_ext if self.input_file else ""
        )
        
        if filename:
            self.output_file = filename
            self.output_entry.delete(0, tk.END)
            self.output_entry.insert(0, filename)
    
    def clear_form(self):
        """Clear the form"""
        self.input_entry.delete(0, tk.END)
        self.output_entry.delete(0, tk.END)
        self.log_text.delete(1.0, tk.END)
        self.input_file = None
        self.output_file = None
    
    def redact_document(self):
        """Redact the selected document"""
        input_file = self.input_entry.get()
        output_file = self.output_entry.get()
        
        if not input_file or not output_file:
            messagebox.showwarning(
                "Missing Information",
                "Please select both input and output files."
            )
            return
        
        if not Path(input_file).exists():
            messagebox.showerror(
                "File Not Found",
                f"Input file not found: {input_file}"
            )
            return
        
        # Disable button and start processing
        self.redact_button.config(state=tk.DISABLED)
        self.progress.start()
        self.status_var.set("Processing...")
        
        # Process in background thread
        threading.Thread(
            target=self.process_document,
            args=(input_file, output_file),
            daemon=True
        ).start()
    
    def process_document(self, input_file, output_file):
        """Process document in background thread"""
        try:
            input_path = Path(input_file)
            file_ext = input_path.suffix.lower()
            
            self.log_message("\n" + "="*70)
            self.log_message(f"Input:  {input_file}")
            self.log_message(f"Output: {output_file}")
            self.log_message("="*70 + "\n")
            
            # Process based on file type
            if file_ext == '.docx':
                stats = self.redactor.redact_docx(input_file, output_file)
                
                self.log_message("\n📊 Redaction Statistics:")
                self.log_message("="*70)
                self.log_message(f"   Paragraphs:  {stats['paragraphs']} redactions")
                self.log_message(f"   Tables:      {stats['tables']} redactions")
                self.log_message(f"   Headers:     {stats['headers']} redactions")
                self.log_message(f"   Footers:     {stats['footers']} redactions")
                self.log_message(f"   {'─'*66}")
                self.log_message(f"   Total:       {stats['total_redactions']} redactions")
                
                if stats['by_type']:
                    self.log_message(f"\n   By Type:")
                    for entity_type, count in sorted(stats['by_type'].items(), 
                                                      key=lambda x: x[1], reverse=True):
                        self.log_message(f"      {entity_type}: {count}")
            
            elif file_ext == '.pdf':
                stats = self.redactor.redact_pdf(input_file, output_file)
                
                self.log_message("\n📊 Redaction Statistics:")
                self.log_message("="*70)
                self.log_message(f"   Pages:       {stats['pages_processed']}")
                self.log_message(f"   Total:       {stats['total_redactions']} redactions")
                
                if stats['by_type']:
                    self.log_message(f"\n   By Type:")
                    for entity_type, count in sorted(stats['by_type'].items(), 
                                                      key=lambda x: x[1], reverse=True):
                        self.log_message(f"      {entity_type}: {count}")
            
            else:
                self.log_message(f"❌ Unsupported file format: {file_ext}")
                self.root.after(0, lambda: messagebox.showerror(
                    "Unsupported Format",
                    f"File format {file_ext} is not supported.\nSupported: .docx, .pdf"
                ))
                return
            
            # Verify redaction
            verification = self.redactor.verify_redaction(output_file, file_ext)
            
            self.log_message("\n" + "="*70)
            if verification.get('safe', False):
                self.log_message("✅ Verification passed - no obvious PII detected in output")
            else:
                self.log_message(f"⚠️  WARNING: Found {verification['pii_found']} potential PII instances!")
                if verification.get('types_found'):
                    self.log_message(f"   Types found:")
                    for entity_type, count in verification['types_found'].items():
                        self.log_message(f"      {entity_type}: {count}")
                self.log_message("   Manual review is CRITICAL!")
            
            self.log_message("="*70)
            self.log_message(f"\n✅ Processing complete!")
            self.log_message(f"   Output saved to: {output_file}")
            
            self.log_message("\n⚠️  IMPORTANT: Manual review is REQUIRED")
            self.log_message("No automated tool is 100% accurate. Please review the output.")
            
            # Show success message
            self.root.after(0, lambda: messagebox.showinfo(
                "Success",
                f"Redaction complete!\n\nOutput saved to:\n{output_file}\n\n"
                "⚠️ Please manually review the redacted document."
            ))
            
        except Exception as e:
            self.log_message(f"\n❌ Error: {e}")
            import traceback
            self.log_message(traceback.format_exc())
            
            self.root.after(0, lambda: messagebox.showerror(
                "Error",
                f"An error occurred during processing:\n\n{str(e)}"
            ))
        
        finally:
            # Re-enable button and stop progress
            self.root.after(0, self.progress.stop)
            self.root.after(0, lambda: self.redact_button.config(state=tk.NORMAL))
            self.root.after(0, lambda: self.status_var.set("Ready"))


def main():
    """Main entry point"""
    try:
        # Check dependencies
        from presidio_analyzer import AnalyzerEngine
        from docx import Document
        import fitz
    except ImportError as e:
        messagebox.showerror(
            "Missing Dependencies",
            f"Required library not found: {e}\n\n"
            "Please install:\n"
            "pip install presidio-analyzer presidio-anonymizer spacy python-docx PyMuPDF\n"
            "python -m spacy download en_core_web_lg"
        )
        sys.exit(1)
    
    root = tk.Tk()
    app = NHSRedactorGUI(root)
    root.mainloop()


if __name__ == "__main__":
    main()