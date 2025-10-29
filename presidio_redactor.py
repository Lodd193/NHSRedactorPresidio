"""
NHS-Specific PII Redactor using Microsoft Presidio
Combines Presidio's ML-powered detection with NHS-specific patterns
"""

import sys
import re
from pathlib import Path
from typing import List, Dict, Optional
from datetime import datetime

try:
    from presidio_analyzer import AnalyzerEngine, Pattern, PatternRecognizer
    from presidio_analyzer.nlp_engine import NlpEngineProvider
    from presidio_anonymizer import AnonymizerEngine
    from presidio_anonymizer.entities import OperatorConfig
except ImportError:
    print("Please install Presidio:")
    print("pip install presidio-analyzer presidio-anonymizer")
    print("pip install spacy")
    print("python -m spacy download en_core_web_lg")
    sys.exit(1)

try:
    from docx import Document
except ImportError:
    print("Please install python-docx: pip install python-docx")
    sys.exit(1)

try:
    import fitz
except ImportError:
    print("Please install PyMuPDF: pip install PyMuPDF")
    sys.exit(1)


class NHSPatternRecognizer(PatternRecognizer):
    """Custom recognizer for NHS-specific PII patterns"""
    
    def __init__(self):
        patterns = [
            Pattern(
                name="nhs_number_pattern",
                regex=r"\b\d{3}[\s-]?\d{3}[\s-]?\d{4}\b",
                score=0.85,
            ),
            Pattern(
                name="hospital_number_pattern",
                regex=r"\b(?:Hospital|Patient|Trust|MRN|Hosp)[\s#:]*([A-Z0-9]{5,12})\b",
                score=0.8,
            ),
            Pattern(
                name="chi_number_pattern",
                regex=r"\bCHI[\s:]*\d{10}\b",
                score=0.9,
            ),
            Pattern(
                name="gmc_number_pattern",
                regex=r"\bGMC[\s:#]*\d{7}\b",
                score=0.9,
            ),
            Pattern(
                name="gp_practice_code_pattern",
                regex=r"\b(?:GP|Practice|Surgery)[\s:]+([A-Z]\d{5})\b",
                score=0.75,
            ),
            Pattern(
                name="uk_driving_license_pattern",
                regex=r"\b[A-Z]{5}\d{6}[A-Z]{2}\d[A-Z]{2}\b",
                score=0.85,
            ),
        ]
        
        super().__init__(
            supported_entity="NHS_IDENTIFIER",
            patterns=patterns,
            context=["NHS", "patient", "hospital", "medical", "clinic", "GP"],
        )


class UKPhoneRecognizer(PatternRecognizer):
    """Custom recognizer for UK phone numbers"""
    
    def __init__(self):
        patterns = [
            Pattern(
                name="uk_mobile_pattern",
                regex=r"\b(?:07\d{3}|447\d{3})[\s-]?\d{6}\b",
                score=0.85,
            ),
            Pattern(
                name="uk_landline_pattern",
                regex=r"\b(?:0|\+44)[\s-]?\d{3,4}[\s-]?\d{3,4}[\s-]?\d{3,4}\b",
                score=0.7,
            ),
        ]
        
        super().__init__(
            supported_entity="UK_PHONE_NUMBER",
            patterns=patterns,
            context=["phone", "mobile", "tel", "telephone", "contact"],
        )


class UKPostcodeRecognizer(PatternRecognizer):
    """Custom recognizer for UK postcodes"""
    
    def __init__(self):
        patterns = [
            Pattern(
                name="uk_postcode_pattern",
                regex=r"\b[A-Z]{1,2}\d{1,2}[A-Z]?\s?\d[A-Z]{2}\b",
                score=0.85,
            ),
        ]
        
        super().__init__(
            supported_entity="UK_POSTCODE",
            patterns=patterns,
            context=["address", "postcode", "postal", "location"],
        )


class UKNINumberRecognizer(PatternRecognizer):
    """Custom recognizer for UK National Insurance Numbers"""
    
    def __init__(self):
        patterns = [
            Pattern(
                name="ni_number_pattern",
                regex=r"\b[A-CEGHJ-PR-TW-Z]{2}[\s-]?\d{2}[\s-]?\d{2}[\s-]?\d{2}[\s-]?[A-D]\b",
                score=0.9,
            ),
        ]
        
        super().__init__(
            supported_entity="UK_NI_NUMBER",
            patterns=patterns,
            context=["NI", "National Insurance", "insurance number"],
        )


class NHSPresidioRedactor:
    """NHS-specific PII redactor using Presidio with custom patterns"""
    
    def __init__(self, language: str = "en"):
        """
        Initialize the NHS Presidio redactor.
        
        Args:
            language: Language code (default: "en")
        """
        # Setup NLP engine with spaCy
        print("🔧 Initializing Presidio analyzer (this may take a moment)...")
        
        configuration = {
            "nlp_engine_name": "spacy",
            "models": [{"lang_code": "en", "model_name": "en_core_web_lg"}],
        }
        
        provider = NlpEngineProvider(nlp_configuration=configuration)
        nlp_engine = provider.create_engine()
        
        # Initialize analyzer with custom recognizers
        self.analyzer = AnalyzerEngine(nlp_engine=nlp_engine)
        
        # Add NHS-specific recognizers
        self.analyzer.registry.add_recognizer(NHSPatternRecognizer())
        self.analyzer.registry.add_recognizer(UKPhoneRecognizer())
        self.analyzer.registry.add_recognizer(UKPostcodeRecognizer())
        self.analyzer.registry.add_recognizer(UKNINumberRecognizer())
        
        # Initialize anonymizer
        self.anonymizer = AnonymizerEngine()
        
        # Define which entities to detect
        self.entities = [
            "PERSON",  # Names (built-in)
            "EMAIL_ADDRESS",  # Emails (built-in)
            "PHONE_NUMBER",  # Generic phones (built-in)
            "DATE_TIME",  # Dates (built-in)
            "LOCATION",  # Locations (built-in)
            "NHS_IDENTIFIER",  # Custom: NHS numbers, hospital IDs, etc.
            "UK_PHONE_NUMBER",  # Custom: UK-specific phones
            "UK_POSTCODE",  # Custom: UK postcodes
            "UK_NI_NUMBER",  # Custom: National Insurance numbers
            "MEDICAL_LICENSE",  # Built-in: Medical licenses
            "UK_NHS",  # Built-in: NHS numbers
        ]
        
        print("✅ Presidio analyzer ready!")
    
    def analyze_text(self, text: str, language: str = "en") -> List:
        """Analyze text and return detected PII entities"""
        results = self.analyzer.analyze(
            text=text,
            language=language,
            entities=self.entities,
            allow_list=None,
            score_threshold=0.3,  # Lower threshold to catch more potential PII
        )
        return results
    
    def anonymize_text(self, text: str, language: str = "en") -> tuple:
        """
        Anonymize text by replacing PII with labels.
        
        Returns:
            tuple: (anonymized_text, number_of_redactions, detailed_findings)
        """
        # Analyze text
        results = self.analyze_text(text, language)
        
        # Create detailed findings
        findings = []
        for result in results:
            findings.append({
                'type': result.entity_type,
                'start': result.start,
                'end': result.end,
                'score': result.score,
                'text': text[result.start:result.end]
            })
        
        # Anonymize with custom labels
        anonymized = self.anonymizer.anonymize(
            text=text,
            analyzer_results=results,
            operators={
                "DEFAULT": OperatorConfig("replace", {"new_value": "[REDACTED]"}),
                "PERSON": OperatorConfig("replace", {"new_value": "[REDACTED-NAME]"}),
                "EMAIL_ADDRESS": OperatorConfig("replace", {"new_value": "[REDACTED-EMAIL]"}),
                "PHONE_NUMBER": OperatorConfig("replace", {"new_value": "[REDACTED-PHONE]"}),
                "UK_PHONE_NUMBER": OperatorConfig("replace", {"new_value": "[REDACTED-UK-PHONE]"}),
                "UK_POSTCODE": OperatorConfig("replace", {"new_value": "[REDACTED-POSTCODE]"}),
                "UK_NI_NUMBER": OperatorConfig("replace", {"new_value": "[REDACTED-NI-NUMBER]"}),
                "NHS_IDENTIFIER": OperatorConfig("replace", {"new_value": "[REDACTED-NHS-ID]"}),
                "UK_NHS": OperatorConfig("replace", {"new_value": "[REDACTED-NHS-NUMBER]"}),
                "DATE_TIME": OperatorConfig("replace", {"new_value": "[REDACTED-DATE]"}),
                "LOCATION": OperatorConfig("replace", {"new_value": "[REDACTED-LOCATION]"}),
                "MEDICAL_LICENSE": OperatorConfig("replace", {"new_value": "[REDACTED-MEDICAL-LICENSE]"}),
            }
        )
        
        return anonymized.text, len(results), findings
    
    def redact_docx(self, input_path: str, output_path: str) -> Dict:
        """
        Redact PII from Word document.
        
        Returns:
            Dictionary with redaction statistics
        """
        print("📄 Loading Word document...")
        doc = Document(input_path)
        
        stats = {
            'paragraphs': 0,
            'tables': 0,
            'headers': 0,
            'footers': 0,
            'total_redactions': 0,
            'by_type': {}
        }
        
        print("🔍 Scanning and redacting paragraphs...")
        # Redact paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                anonymized, count, findings = self.anonymize_text(para.text)
                if count > 0:
                    para.text = anonymized
                    stats['paragraphs'] += count
                    self._update_type_stats(stats, findings)
        
        print("🔍 Scanning and redacting tables...")
        # Redact tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        anonymized, count, findings = self.anonymize_text(cell.text)
                        if count > 0:
                            cell.text = anonymized
                            stats['tables'] += count
                            self._update_type_stats(stats, findings)
        
        print("🔍 Scanning and redacting headers/footers...")
        # Redact headers and footers
        for section in doc.sections:
            for header in section.header.paragraphs:
                if header.text.strip():
                    anonymized, count, findings = self.anonymize_text(header.text)
                    if count > 0:
                        header.text = anonymized
                        stats['headers'] += count
                        self._update_type_stats(stats, findings)
            
            for footer in section.footer.paragraphs:
                if footer.text.strip():
                    anonymized, count, findings = self.anonymize_text(footer.text)
                    if count > 0:
                        footer.text = anonymized
                        stats['footers'] += count
                        self._update_type_stats(stats, findings)
        
        stats['total_redactions'] = (stats['paragraphs'] + stats['tables'] + 
                                     stats['headers'] + stats['footers'])
        
        print("💾 Saving redacted document...")
        doc.save(output_path)
        
        return stats
    
    def redact_pdf(self, input_path: str, output_path: str) -> Dict:
        """
        Redact PII from PDF document.
        
        Returns:
            Dictionary with redaction statistics
        """
        print("📄 Loading PDF document...")
        doc = fitz.open(input_path)
        
        stats = {
            'pages_processed': 0,
            'total_redactions': 0,
            'by_type': {}
        }
        
        print(f"🔍 Scanning and redacting {len(doc)} pages...")
        for page_num, page in enumerate(doc, 1):
            # Extract text
            text = page.get_text()
            
            # Analyze for PII
            results = self.analyze_text(text)
            
            if results:
                for result in results:
                    # Find the text on the page
                    pii_text = text[result.start:result.end]
                    areas = page.search_for(pii_text)
                    
                    if areas:
                        for area in areas:
                            page.add_redact_annot(area, fill=(0, 0, 0))
                            stats['total_redactions'] += 1
                            
                            # Update type stats
                            entity_type = result.entity_type
                            stats['by_type'][entity_type] = stats['by_type'].get(entity_type, 0) + 1
                
                # Apply redactions
                page.apply_redactions()
            
            stats['pages_processed'] += 1
            
            if page_num % 10 == 0:
                print(f"   Processed {page_num}/{len(doc)} pages...")
        
        print("💾 Saving redacted PDF...")
        doc.save(output_path)
        doc.close()
        
        return stats
    
    def _update_type_stats(self, stats: Dict, findings: List[Dict]):
        """Update statistics by PII type"""
        for finding in findings:
            entity_type = finding['type']
            stats['by_type'][entity_type] = stats['by_type'].get(entity_type, 0) + 1
    
    def verify_redaction(self, file_path: str, file_type: str) -> Dict:
        """
        Verify that PII has been removed from the redacted file.
        
        Returns:
            Dictionary with verification results
        """
        print("🔍 Verifying redaction...")
        verification = {
            'pii_found': 0,
            'types_found': {},
            'safe': True
        }
        
        try:
            if file_type == '.docx':
                doc = Document(file_path)
                full_text = '\n'.join([para.text for para in doc.paragraphs])
                
                # Also check tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            full_text += '\n' + cell.text
            
            elif file_type == '.pdf':
                doc = fitz.open(file_path)
                full_text = ''
                for page in doc:
                    full_text += page.get_text()
                doc.close()
            
            else:
                return verification
            
            # Scan for any remaining PII
            results = self.analyze_text(full_text)
            
            if results:
                verification['safe'] = False
                verification['pii_found'] = len(results)
                
                for result in results:
                    entity_type = result.entity_type
                    verification['types_found'][entity_type] = \
                        verification['types_found'].get(entity_type, 0) + 1
        
        except Exception as e:
            verification['error'] = str(e)
        
        return verification


def main():
    if len(sys.argv) < 2:
        print("=" * 70)
        print("NHS PII Redactor - Presidio Edition")
        print("Powered by Microsoft Presidio + NHS-Specific Patterns")
        print("=" * 70)
        print("\nUsage: python nhs_presidio_redactor.py <input_file> [output_file]")
        print("\nSupported formats: .docx, .pdf")
        print("\nFeatures:")
        print("  ✓ ML-powered name detection (spaCy)")
        print("  ✓ NHS-specific patterns (NHS numbers, CHI, GMC, etc.)")
        print("  ✓ UK-specific patterns (NI numbers, postcodes, phones)")
        print("  ✓ Context-aware detection")
        print("  ✓ Automatic verification")
        print("\nDetects:")
        print("  • Names (ML-powered)")
        print("  • NHS Numbers & Hospital IDs")
        print("  • Email Addresses")
        print("  • UK Phone Numbers")
        print("  • UK Postcodes")
        print("  • National Insurance Numbers")
        print("  • Dates & Locations")
        print("  • Medical Licenses")
        print("  • And more...")
        print("\n" + "=" * 70)
        print("\nFirst-time setup required:")
        print("  pip install presidio-analyzer presidio-anonymizer spacy")
        print("  python -m spacy download en_core_web_lg")
        print("=" * 70 + "\n")
        sys.exit(1)
    
    input_file = Path(sys.argv[1])
    
    if not input_file.exists():
        print(f"❌ Error: File '{input_file}' not found")
        sys.exit(1)
    
    # Determine output file
    if len(sys.argv) > 2:
        output_file = Path(sys.argv[2])
    else:
        output_file = input_file.parent / f"{input_file.stem}_REDACTED{input_file.suffix}"
    
    print(f"\n🔒 NHS PII Redactor - Presidio Edition")
    print(f"{'=' * 70}")
    print(f"Input:  {input_file}")
    print(f"Output: {output_file}")
    print(f"{'=' * 70}\n")
    
    try:
        # Initialize redactor
        redactor = NHSPresidioRedactor()
        
        # Process file
        if input_file.suffix.lower() == '.docx':
            stats = redactor.redact_docx(str(input_file), str(output_file))
            
            print(f"\n📊 Redaction Statistics:")
            print(f"{'=' * 70}")
            print(f"   Paragraphs:  {stats['paragraphs']} redactions")
            print(f"   Tables:      {stats['tables']} redactions")
            print(f"   Headers:     {stats['headers']} redactions")
            print(f"   Footers:     {stats['footers']} redactions")
            print(f"   {'─' * 66}")
            print(f"   Total:       {stats['total_redactions']} redactions")
            
            if stats['by_type']:
                print(f"\n   By Type:")
                for entity_type, count in sorted(stats['by_type'].items(), 
                                                  key=lambda x: x[1], reverse=True):
                    print(f"      {entity_type}: {count}")
            
        elif input_file.suffix.lower() == '.pdf':
            stats = redactor.redact_pdf(str(input_file), str(output_file))
            
            print(f"\n📊 Redaction Statistics:")
            print(f"{'=' * 70}")
            print(f"   Pages:       {stats['pages_processed']}")
            print(f"   Total:       {stats['total_redactions']} redactions")
            
            if stats['by_type']:
                print(f"\n   By Type:")
                for entity_type, count in sorted(stats['by_type'].items(), 
                                                  key=lambda x: x[1], reverse=True):
                    print(f"      {entity_type}: {count}")
        
        else:
            print(f"❌ Error: Unsupported file format '{input_file.suffix}'")
            print("Supported formats: .docx, .pdf")
            sys.exit(1)
        
        # Verify redaction
        verification = redactor.verify_redaction(str(output_file), 
                                                 input_file.suffix.lower())
        
        print(f"\n{'=' * 70}")
        if verification.get('safe', False):
            print("✅ Verification passed - no obvious PII detected in output")
        else:
            print(f"⚠️  WARNING: Found {verification['pii_found']} potential PII instances!")
            if verification.get('types_found'):
                print(f"   Types found:")
                for entity_type, count in verification['types_found'].items():
                    print(f"      {entity_type}: {count}")
            print("   Manual review is CRITICAL!")
        
        print(f"{'=' * 70}")
        print(f"\n✅ Processing complete!")
        print(f"   Output saved to: {output_file}")
        
        print("\n" + "=" * 70)
        print("⚠️  CRITICAL: Manual Review Required")
        print("=" * 70)
        print("While Presidio uses advanced ML, no tool is 100% accurate.")
        print("You MUST manually review the redacted document for:")
        print("  • Names in unusual contexts")
        print("  • Addresses and locations")
        print("  • Sensitive contextual information")
        print("  • Document metadata and properties")
        print("\nNOTE: The following are NOT automatically redacted:")
        print("  • Text boxes and shapes (DOCX)")
        print("  • Comments and tracked changes (DOCX)")
        print("  • Document properties and metadata")
        print("  • Text within images")
        print("=" * 70 + "\n")
    
    except ImportError as e:
        print(f"\n❌ Missing dependency: {e}")
        print("\nPlease install required packages:")
        print("  pip install presidio-analyzer presidio-anonymizer spacy")
        print("  python -m spacy download en_core_web_lg")
        sys.exit(1)
    
    except Exception as e:
        print(f"\n❌ Error processing file: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()